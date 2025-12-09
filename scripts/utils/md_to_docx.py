import sys
import re
import os

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Error: python-docx is not installed.")
    print("Please run: pip install python-docx")
    sys.exit(1)

def parse_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: File not found: {md_path}")
        return

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    print(f"Reading {md_path}...")
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    table_mode = False
    table_lines = []
    code_mode = False

    for line in lines:
        line = line.rstrip()
        
        # Code Blocks
        if line.startswith('```'):
            code_mode = not code_mode
            continue
        if code_mode:
            try:
                p = doc.add_paragraph(line)
                p.style = 'No Spacing'
                if p.runs:
                    p.runs[0].font.name = 'Courier New'
            except:
                pass # Fallback if style fails
            continue

        # Tables
        if line.startswith('|'):
            table_mode = True
            table_lines.append(line)
            continue
        elif table_mode:
            # End of table detected (line doesn't start with |)
            if len(table_lines) > 2: # Header, separator, body...
                try:
                    # Basic table parsing
                    headers = [c.strip() for c in table_lines[0].strip('|').split('|')]
                    rows = [l.strip('|').split('|') for l in table_lines[2:]] # Skip separator
                    
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Table Grid'
                    
                    hdr_cells = table.rows[0].cells
                    for i, h in enumerate(headers):
                        if i < len(hdr_cells):
                            hdr_cells[i].text = h
                            hdr_cells[i].paragraphs[0].runs[0].bold = True
                        
                    for row_data in rows:
                        row_cells = table.add_row().cells
                        for i, cell_data in enumerate(row_data):
                            if i < len(row_cells):
                                row_cells[i].text = cell_data.strip()
                except Exception as e:
                    print(f"Warning: Failed to parse table: {e}")
                    # Fallback: dump as text
                    for tl in table_lines:
                        doc.add_paragraph(tl, style='No Spacing')
            
            table_mode = False
            table_lines = []
            
            if not line: continue

        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Lists
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            doc.add_paragraph(line.split('.', 1)[1].strip(), style='List Number')
            
        # Normal Text
        else:
            if line:
                p = doc.add_paragraph()
                # Basic bold handling **text**
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

    doc.save(docx_path)
    print(f"Success! Created {docx_path}")

if __name__ == "__main__":
    parse_md_to_docx('PROJECT_DOCUMENTATION.md', 'PROJECT_DOCUMENTATION.docx')
